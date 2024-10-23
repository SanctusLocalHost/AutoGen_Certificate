import os
import tkinter as tk
from tkinter import messagebox
from tkinterdnd2 import TkinterDnD, DND_FILES
from tkinter import ttk
from docx import Document
from docx.shared import Pt

# Certifique-se de que o pywin32 esteja instalado:
try:
    import win32com.client as win32
except ImportError:
    messagebox.showerror("Erro", "O módulo 'pywin32' não está instalado. Execute 'pip install pywin32' para instalá-lo.")
    raise

# Dicionário de substituição EAN para códigos NX
substituicoes_ean_para_nx = {
    "7898740470008,1": "NX 028-ABRAC",
    "7898740470015,1": "NX 001",
    "7898740470022,1": "NX 002",
    "7898740470039,1": "NX 003",
    "7898740470046,1": "NX 004",
    "7898740470053,1": "NX 005",
    "7898740470060,1": "NX 006",
    "7898740470077,1": "NX 007",
    "7898740470084,1": "NX 008",
    "7898740470091,1": "NX 009",
    "7898740470107,1": "NX 010",
    "7898740470114,1": "NX 011",
    "7898740470121,1": "NX 012",
    "7898740470138,1": "NX 013",
    "7898740470145,1": "NX 014",
    "7898740470152,1": "NX 015",
    "7898740470169,1": "NX 016",
    "7898740470176,1": "NX 017",
    "7898740470183,1": "NX 018",
    "7898740470190,1": "NX 019",
    "7898740470206,1": "NX 020",
    "7898740470213,1": "NX 021",
    "7898740470220,1": "NX 022",
    "7898740470237,1": "NX 023",
    "7898740470244,1": "NX 024",
    "7898740470251,1": "NX 025",
    "7898740470268,1": "NX 026",
    "7898740470275,1": "NX 027",
    "7898740470282,1": "NX 028",
    "7898740470299,1": "NX 029",
    "7898740470305,1": "NX 030",
    "7898740470312,1": "NX 031",
    "7898740470329,1": "NX 032",
    "7898740470336,1": "NX 033",
    "7898740470343,1": "NX 034",
    "7898740470350,1": "NX 035",
    "7898740470367,1": "NX 036",
    "7898740470374,1": "NX 037",
    "7898740470381,1": "NX 038",
    "7898740470398,1": "NX 039",
    "7898740470404,1": "NX 040",
    "7898740470411,1": "NX 041",
    "7898740470428,1": "NX 042",
    "7898740470435,1": "NX 043",
    "7898740470442,1": "NX 044",
    "7898740470459,1": "NX 045",
    "7898740470466,1": "NX 046",
    "7898740470473,1": "NX 047",
    "7898740470480,1": "NX 048",
    "7898740470497,1": "NX 049",
    "7898740470503,1": "NX 050",
    "7898740470510,1": "NX 051",
    "7898740470527,1": "NX 052",
    "7898740470534,1": "NX 053",
    "7898740470541,1": "NX 054",
    "7898740470558,1": "NX 055",
    "7898740470565,1": "NX 056",
    "7898740470572,1": "NX 057",
    "7898740470589,1": "NX 058",
    "7898740470596,1": "NX 059",
    "7898740470602,1": "NX 060",
    "7898740470619,1": "NX 061",
    "7898740470626,1": "NX 062",
    "7898740470633,1": "NX 063",
    "7898740470640,1": "NX 064",
    "7898740470657,1": "NX 065",
    "7898740470664,1": "NX 066",
    "7898740470671,1": "NX 067",
    "7898740470688,1": "NX 068",
    "7898740470695,1": "NX 069",
    "7898740470701,1": "NX 070",
    "7898740470718,1": "NX 071",
    "7898740470725,1": "NX 072",
    "7898740470732,1": "NX 073",
    "7898740470749,1": "NX 074",
    "7898740470756,1": "NX 075",
    "7898740470763,1": "NX 076",
    "7898740470770,1": "NX 077",
    "7898740470787,1": "NX 078",
    "7898740470794,1": "NX 079",
    "7898740470800,1": "NX 080",
    "7898740470817,1": "NX 081",
    "7898740470824,1": "NX 082",
    "7898740470831,1": "NX 083",
    "7898740470848,1": "NX 084",
    "7898740470855,1": "NX 085",
    "7898740470862,1": "NX 086",
    "7898740470879,1": "NX 087",
    "7898740470886,1": "NX 088",
    "7898740470893,1": "NX 089",
    "7898740470909,1": "NX 090",
    "7898740470916,1": "NX 091",
    "7898740470923,1": "NX 092",
    "7898740470930,1": "NX 093",
    "7898740470947,1": "NX 094",
    "7898740470954,1": "NX 095",
    "7898740470961,1": "NX 096",
    "7898740470978,1": "NX 097",
    "7898740470985,1": "NX 098",
    "7898740470992,1": "NX 099",
    "7898740471005,1": "NX 100",
    "7898740471012,1": "NX 101",
    "7898740471029,1": "NX 102",
    "7898740471036,1": "NX 103",
    "7898740471043,1": "NX 104",
    "7898740471050,1": "NX 105",
    "7898740471067,1": "NX 106",
    "7898740471074,1": "NX 107",
    "7898740471081,1": "NX 108",
    "7898740471098,1": "NX 109",
    "7898740471104,1": "NX 110",
    "7898740471111,1": "NX 111",
    "7898740471128,1": "NX 112",
    "7898740471135,1": "NX 113",
    "7898740471142,1": "NX 114",
    "7898740471159,1": "NX 115",
    "7898740471166,1": "NX 116",
    "7898740471173,1": "NX 117",
    "7898740471180,1": "NX 118",
    "7898740471197,1": "NX 119",
    "7898740471203,1": "NX 120",
    "7898740471210,1": "NX 121",
    "7898740471227,1": "NX 122",
    "7898740471234,1": "NX 123",
    "7898740471241,1": "NX 124",
    "7898740471258,1": "NX 125",
    "7898740471265,1": "NX 126",
    "7898740471272,1": "NX 127",
    "7898740471289,1": "NX 128",
    "7898740471296,1": "NX 129",
    "7898740471302,1": "NX 130",
    "7898740471319,1": "NX 131",
    "7898740471326,1": "NX 132",
    "7898740471333,1": "NX 133",
    "7898740471340,1": "NX 134",
    "7898740471357,1": "NX 135",
    "7898740471364,1": "NX 136",
    "7898740471371,1": "NX 137",
    "7898740471388,1": "NX 138",
    "7898740471395,1": "NX 139",
    "7898740471401,1": "NX 140",
    "7898740471418,1": "NX 141",
    "7898740471425,1": "NX 142",
    "7898740471432,1": "NX 143",
    "7898740471449,1": "NX 144",
    "7898740471456,1": "NX 145",
    "7898740471463,1": "NX 146",
    "7898740471470,1": "NX 147",
    "7898740471487,1": "NX 148",
    "7898740471494,1": "NX 149",
    "7898740471500,1": "NX 150",
    "7898740471517,1": "NX 151",
    "7898740471524,1": "NX 152",
    "7898740471531,1": "NX 153",
    "7898740471548,1": "NX 154",
    "7898740471555,1": "NX 155",
    "7898740471562,1": "NX 156",
    "7898740471579,1": "NX 157",
    "7898740471586,1": "NX 158",
    "7898740471593,1": "NX 159",
    "7898740471609,1": "NX 160",
    "7898740471616,1": "NX 161",
    "7898740471623,1": "NX 162",
    "7898740471630,1": "NX 163",
    "7898740471647,1": "NX 164",
    "7898740471654,1": "NX 165",
    "7898740471661,1": "NX 166",
    "7898740471678,1": "NX 167",
    "7898740471685,1": "NX 168",
    "7898740471692,1": "NX 169",
    "7898740471708,1": "NX 170",
    "7898740471715,1": "NX 171",
    "7898740471722,1": "NX 172",
    "7898740471739,1": "NX 173",
    "7898740471746,1": "NX 174",
    "7898740471753,1": "NX 175",
    "7898740471760,1": "NX 176",
    "7898740471777,1": "NX 177",
    "7898740471784,1": "NX 178",
    "7898740471791,1": "NX 179",
    "7898740471807,1": "NX 180",
    "7898740471814,1": "NX 181",
    "7898740471821,1": "NX 182",
    "7898740471838,1": "NX 183",
    "7898740471845,1": "NX 184",
    "7898740471852,1": "NX 185",
    "7898740471869,1": "NX 186",
    "7898740471876,1": "NX 187",
    "7898740471883,1": "NX 188",
    "7898740471890,1": "NX 189",
    "7898740471906,1": "NX 190",
    "7898740471913,1": "NX 191",
    "7898740471920,1": "NX 192",
    "7898740471937,1": "NX 193",
    "7898740471944,1": "NX 194",
    "7898740471951,1": "NX 195",
    "7898740471968,1": "NX 196",
    "7898740471975,1": "NX 197",
    "7898740471982,1": "NX 198",
    "7898740471999,1": "NX 199",
    "7898740472002,1": "NX 200",
    "7898740472019,1": "NX 201",
    "7898740472026,1": "NX 202",
    "7898740472033,1": "NX 203",
    "7898740472040,1": "NX 204",
    "7898740472057,1": "NX 205",
    "7898740472064,1": "NX 206",
    "7898740472071,1": "NX 207",
    "7898740472088,1": "NX 208",
    "7898740472095,1": "NX 209",
    "7898740472101,1": "NX 210",
    "7898740472118,1": "NX 211",
    "7898740472125,1": "NX 212",
    "7898740472132,1": "NX 213",
    "7898740472149,1": "NX 214",
    "7898740472156,1": "NX 215",
    "7898740472163,1": "NX 216",
    "7898740472170,1": "NX 217",
    "7898740472187,1": "NX 218",
    "7898740472194,1": "NX 219",
    "7898740472200,1": "NX 220",
    "7898740472217,1": "NX 221",
    "7898740472224,1": "NX 222",
    "7898740472231,1": "NX 223",
    "7898740472248,1": "NX 224",
    "7898740472255,1": "NX 225",
    "7898740472262,1": "NX 226",
    "7898740472279,1": "NX 227",
    "7898740472286,1": "NX 228",
    "7898740472293,1": "NX 229",
    "7898740472309,1": "NX 230",
    "7898740472316,1": "NX 231",
    "7898740472323,1": "NX 232",
    "7898740472330,1": "NX 233",
    "7898740472347,1": "NX 234",
    "7898740472354,1": "NX 235",
    "7898740472361,1": "NX 236",
    "7898740472378,1": "NX 237",
    "7898740472385,1": "NX 238",
    "7898740472392,1": "NX 239",
    "7898740472408,1": "NX 240",
    "7898740472415,1": "NX 241",
    "7898740472422,1": "NX 242",
    "7898740472439,1": "NX 243",
    "7898740472446,1": "NX 244",
    "7898740472453,1": "NX 245",
    "7898740472460,1": "NX 246",
    "7898740472477,1": "NX 247",
    "7898740472484,1": "NX 248",
    "7898740472491,1": "NX 249",
    "7898740472507,1": "NX 250",
    "7898740472514,1": "NX 251",
    "7898740472521,1": "NX 252",
    "7898740472538,1": "NX 253",
	"7898740472545,1": "NX 254",
    "7898740472552,1": "NX 255",
    "7898740472569,1": "NX 256",
    "7898740472576,1": "NX 257",
    "7898740472583,1": "NX 258",
	"7894325102432,1": "523402",
    "7894325103644,1": "523437",
    "7894325000110,1": "526033",
    "7894325000271,1": "526034",
    "7894325000264,1": "526036",
    "7894325008147,1": "526044",
    "7894325000943,1": "526048",
    "7894325002510,1": "526055",
    "7894325003258,1": "526059",
    "7894325004156,1": "526065",
    "7894325002848,1": "526198",
    "7894325004071,1": "526199",
    "7894325004064,1": "526200",
    "7894325000066,1": "526201",
    "7894325002183,1": "526202",
    "7894325000059,1": "526203",
    "7894325008499,1": "526205",
    "7894325001667,1": "526206",
    "7894325001247,1": "526211",
    "7894325003968,1": "526212",
    "7894325001650,1": "526291",
    "7894325001674,1": "526292",
    "7894325003777,1": "526294",
    "7894325000837,1": "526297",
    "7894325002138,1": "526388",
    "7894325002565,1": "526389",
    "7894325001513,1": "526391",
    "7894325004514,1": "526392",
    "7894325001940,1": "526393",
    "7894325003241,1": "526399",
    "7894325001940,1": "526488",
    "7894325100001,1": "526489",
    "7894325001261,1": "526502",
    "7894325001605,1": "526503",
    "7894325003807,1": "526504",
    "7894325004606,1": "526505",
    "7894325001728,1": "526506",
    "7894325001438,1": "526507",
    "7894325002206,1": "526508",
    "7894325004163,1": "526509",
    "7894325004972,1": "526511",
    "7894325000295,1": "526538",
    "7894325001285,1": "526539",
    "7894325004989,1": "526540",
    "7894325002435,1": "526541",
    "7894325005276,1": "526542",
    "7894325100988,1": "523391D",
    "7894325100988,1": "523391E",
}

def convert_files(files):
    # Criar a pasta "BIPAGEM CONVERTIDA" no diretório atual do script
    script_directory = os.path.dirname(os.path.abspath(__file__))
    converted_folder = os.path.join(script_directory, "BIPAGEM CONVERTIDA")
    os.makedirs(converted_folder, exist_ok=True)

    converted_files = []

    # Processar cada arquivo e converter EAN para NX se necessário
    for file_path in files:
        converted_file_path = os.path.join(converted_folder, os.path.basename(file_path))
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao ler o arquivo {file_path}:\n{e}")
            continue

        try:
            with open(converted_file_path, 'w', encoding='utf-8') as converted_file:
                for line in lines:
                    parts = line.strip().split()
                    if len(parts) == 2:
                        code = parts[0]
                        quantity = parts[1]
                        # Verificar se é EAN ou NX
                        if code.startswith("NX"):
                            item = code
                        else:
                            # Usar o código diretamente como chave
                            item = substituicoes_ean_para_nx.get(code, code)
                        converted_file.write(f"{item} {quantity}\n")
                    else:
                        converted_file.write(line)
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao escrever o arquivo convertido {converted_file_path}:\n{e}")
            continue

        converted_files.append(converted_file_path)

    return converted_files

def process_files(destination, order_number, files):
    # Caminho base onde as pastas dos clientes estão localizadas
    base_path = r'G:\Meu Drive\CONTROLLER\MINUTA'

    # Construir o caminho da pasta do cliente
    client_folder = os.path.join(base_path, destination)

    # Verificar se a pasta do cliente existe
    if not os.path.exists(client_folder):
        messagebox.showerror("Erro", f"A pasta do cliente '{destination}' não foi encontrada em '{base_path}'.")
        return

    # Carregar o template do Word
    script_directory = os.path.dirname(os.path.abspath(__file__))
    template_path = r"G:\Meu Drive\CONTROLLER\MINUTA\MODELO_MINUTA_TRANSPORTE_AUTOMATICA.docx"

    if not os.path.exists(template_path):
        messagebox.showerror("Erro", f"Template não encontrado: {template_path}")
        return

    # Carregar o documento
    document = Document(template_path)

    # Substituir as variáveis de destino e pedido e aplicar negrito
    for paragraph in document.paragraphs:
        if 'DESTINO:' in paragraph.text:
            paragraph.text = paragraph.text.replace('DESTINO:', f'DESTINO: {destination}')
            paragraph.runs[0].font.bold = True
        if 'PEDIDO:' in paragraph.text:
            paragraph.text = paragraph.text.replace('PEDIDO:', f'PEDIDO: {order_number}')
            paragraph.runs[0].font.bold = True

    # Processar cada arquivo de texto e preencher as tabelas no Word
    for file_path in sorted(files):
        # Extrair o nome do arquivo sem extensão
        file_name = os.path.basename(file_path)
        file_name_no_ext = os.path.splitext(file_name)[0]

        # Determinar o título com base no nome do arquivo
        lower_file_name = file_name_no_ext.lower()
        if 'saco' in lower_file_name or 'volume' in lower_file_name:
            # Extrair número se presente
            words = file_name_no_ext.split()
            number = words[-1] if words[-1].isdigit() else ''
            section_title = f"SACO {number}" if number else "SACO"
        elif 'palete' in lower_file_name or 'pallet' in lower_file_name:
            # Extrair número se presente
            words = file_name_no_ext.split()
            number = words[-1] if words[-1].isdigit() else ''
            section_title = f"PALETE {number}" if number else "PALETE"
        else:
            # Se não encontrar, usar "SACO" por padrão
            section_title = f"SACO"

        # Adicionar título da seção com negrito e tamanho 14
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(section_title)
        title_run.font.bold = True
        title_run.font.size = Pt(14)

        # Criar tabela para a seção
        table = document.add_table(rows=2, cols=2)
        table.style = 'Table Grid'

        # Cabeçalhos da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DESCRIÇÃO'
        hdr_cells[1].text = 'QUANTIDADE'

        # Alinhar cabeçalhos à esquerda e em negrito
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.bold = True
                paragraph.alignment = 0  # Alinhar à esquerda

        # Inicializar as células de descrição e quantidade
        description_text = ""
        quantity_text = ""

        # Ler o conteúdo do arquivo e adicionar à célula de descrição e quantidade
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
                for line in lines:
                    parts = line.strip().split()
                    if len(parts) >= 2:
                        quantity = parts[-1]
                        item = ' '.join(parts[:-1])
                    else:
                        continue

                    # Adicionar cada item e quantidade em uma nova linha dentro da célula
                    description_text += f"{item}\n"
                    quantity_text += f"{quantity}\n"
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao ler o arquivo convertido {file_path}:\n{e}")
            continue

        # Atribuir o texto concatenado às células da segunda linha
        table.rows[1].cells[0].text = description_text.strip()
        table.rows[1].cells[1].text = quantity_text.strip()

        # Alinhar o texto à esquerda nas células de descrição e quantidade
        for cell in table.rows[1].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 0  # Alinhar à esquerda
                for run in paragraph.runs:
                    run.font.size = Pt(12)

        # Adicionar um espaço após cada tabela de seção
        document.add_paragraph()

    # Salvar o arquivo na pasta do cliente
    docx_filename = f"MINUTA {destination} {order_number}.docx"
    pdf_filename = f"MINUTA {destination} {order_number}.pdf"
    docx_path = os.path.join(client_folder, docx_filename)
    pdf_path = os.path.join(client_folder, pdf_filename)

    try:
        document.save(docx_path)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao salvar o arquivo Word:\n{e}")
        return

    # Converter para PDF
    try:
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        word_doc = word.Documents.Open(docx_path)
        word_doc.SaveAs(pdf_path, FileFormat=17)  # 17 é o formato PDF
        word_doc.Close()
        word.Quit()
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao converter para PDF:\n{e}")
        return

    # Remover o arquivo Word temporário
    try:
        os.remove(docx_path)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao remover o arquivo Word temporário:\n{e}")

    messagebox.showinfo("SUCESSO!", f"Arquivo PDF salvo em:\n{pdf_path}")

def generate_document():
    destination = destination_entry.get()
    order_number = order_entry.get()
    files = file_listbox.get(0, tk.END)

    if not destination or not order_number or not files:
        messagebox.showerror("Erro", "Por favor, preencha todos os campos e selecione os arquivos.")
        return

    converted_files = convert_files(files)
    if not converted_files:
        messagebox.showerror("Erro", "Nenhum arquivo foi convertido.")
        return
    process_files(destination, order_number, converted_files)

dragged_files = []

dragged_files = []


def reset_fields():
    # Limpar os campos de entrada
    destination_entry.delete(0, tk.END)
    order_entry.delete(0, tk.END)
    file_listbox.delete(0, tk.END)

    # Remover arquivos da pasta "BIPAGEM CONVERTIDA"
    script_directory = os.path.dirname(os.path.abspath(__file__))
    converted_folder = os.path.join(script_directory, "BIPAGEM CONVERTIDA")
    if os.path.exists(converted_folder):
        for filename in os.listdir(converted_folder):
            file_path = os.path.join(converted_folder, filename)
            try:
                os.remove(file_path)
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao deletar o arquivo {file_path}: {e}")

        # Tentar remover a pasta "BIPAGEM CONVERTIDA" após deletar os arquivos
        try:
            os.rmdir(converted_folder)  # Remove a pasta, já que está vazia agora
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao deletar a pasta {converted_folder}: {e}")

    # Deletar os arquivos arrastados
    global dragged_files
    for file_path in dragged_files:
        try:
            os.remove(file_path)
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao deletar o arquivo {file_path}: {e}")

    # Resetar a lista de arquivos arrastados
    dragged_files = []

def generate_excel():
    destination = destination_entry.get()
    order_number = order_entry.get()
    files = file_listbox.get(0, tk.END)

    if not destination or not order_number or not files:
        messagebox.showerror("Erro", "Por favor, preencha todos os campos e selecione os arquivos.")
        return

    process_files(destination, order_number, files)

def drop(event):
    files = root.tk.splitlist(event.data)
    for file in files:
        file_listbox.insert(tk.END, file)
        dragged_files.append(file)

# Interface gráfica
root = TkinterDnD.Tk()
root.title("AUTOGEN MINUTAS ULTIMATE V1.9")

style = ttk.Style(root)
style.configure("TLabel", font=("Helvetica", 12))
style.configure("TButton", font=("Helvetica", 8), padding=5)
style.configure("TEntry", font=("Helvetica", 12))

ttk.Label(root, text="Nome do Cliente:").grid(row=0, column=0, padx=10, pady=5)
destination_entry = ttk.Entry(root)
destination_entry.grid(row=0, column=1, padx=10, pady=5)

ttk.Label(root, text="Nº Doc.Saída ou NF:").grid(row=1, column=0, padx=10, pady=5)
order_entry = ttk.Entry(root)
order_entry.grid(row=1, column=1, padx=10, pady=5)

ttk.Label(root, text="Arraste os arquivos da bipagem e clique em Gerar Minuta",
          font=("Helvetica", 10, "italic")).grid(row=2, column=0, columnspan=2, padx=10, pady=5)

file_listbox = tk.Listbox(root, selectmode=tk.MULTIPLE, width=50, height=10, bg="white", fg="black",
                          font=("Helvetica", 10))
file_listbox.grid(row=3, column=0, columnspan=2, padx=10, pady=5)

# Botões
generate_button = ttk.Button(root, text="Gerar Minuta", command=generate_document)
generate_button.grid(row=4, column=0, padx=2, pady=10)

reset_button = ttk.Button(root, text="Redefinir", command=reset_fields)
reset_button.grid(row=4, column=1, padx=2, pady=10)

file_listbox.drop_target_register(DND_FILES)
file_listbox.dnd_bind('<<Drop>>', drop)

root.mainloop()
